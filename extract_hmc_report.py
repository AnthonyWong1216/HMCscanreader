#!/usr/bin/env python3
"""
Script to extract HMC and System Summary information from Excel files
and generate a Word document report with tables.
"""

import os
import glob
import re
import json
from collections import defaultdict
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import openpyxl

def clean_text(text):
    """Clean text by removing escape characters and ANSI codes."""
    if not text:
        return ""
    # Convert to string first to handle different data types
    text = str(text)
    # Remove ANSI escape sequences
    ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')
    text = ansi_escape.sub('', text)
    # Remove other control characters
    text = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', text)
    return text.strip()

def extract_data_from_excel(file_path):
    """Extract data from Excel file."""
    data = {
        'hmc_info': [],
        'system_info': [],
        'lpar_info': []
    }
    
    try:
        # Try to read with pandas first
        try:
            # Read all sheets from the Excel file
            excel_file = pd.ExcelFile(file_path)
            print(f"Found sheets: {excel_file.sheet_names}")
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                print(f"Processing sheet: {sheet_name}")
                print(f"Columns: {list(df.columns)}")
                print(f"Shape: {df.shape}")
                
                # Process the data based on sheet content
                process_excel_sheet(df, sheet_name, data)
                
        except Exception as e:
            print(f"Pandas failed, trying openpyxl: {e}")
            # Fallback to openpyxl
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                print(f"Processing sheet: {sheet_name}")
                
                # Convert sheet to DataFrame-like structure
                rows = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        rows.append(row)
                
                if rows:
                    # Create a simple DataFrame-like structure
                    df = pd.DataFrame(rows[1:], columns=rows[0] if rows else [])
                    process_excel_sheet(df, sheet_name, data)
                    
    except Exception as e:
        print(f"Error reading Excel file {file_path}: {e}")
    
    return data

def process_excel_sheet(df, sheet_name, data):
    """Process individual Excel sheet and extract relevant data."""
    if df.empty:
        return
    
    # Convert DataFrame to string for easier searching
    df_str = df.astype(str)
    
    # Look for HMC information in specific tab
    if 'hmc' in sheet_name.lower():
        extract_hmc_info_from_sheet(df, data)
    
    # Look for System information in specific tab
    elif 'system_summary' in sheet_name.lower():
        extract_system_summary_info_from_sheet(df, data)
    
    # Look for LPAR information in specific tab
    elif 'lpar_profiles' in sheet_name.lower():
        extract_lpar_info_from_sheet(df, data)

def extract_hmc_info_from_sheet(df, data):
    """Extract HMC information from Excel sheet using specific cell locations."""
    try:
        print(f"Extracting HMC info from sheet with columns: {list(df.columns)}")
        print(f"DataFrame shape: {df.shape}")
        
        hmc_info = {
            'hostname': None,
            'ip_addresses': {},
            'hardware_model': None,
            'serial': None,
            'base_version': None,
            'service_pack': None,
            'gateway': None
        }
        
        # Get values from specific cell locations
        try:
            # Hostname: B17 (row 15, column 1) - Excel row 17 = pandas row 15
            if len(df) > 15 and len(df.columns) > 1:
                hmc_info['hostname'] = clean_text(df.iloc[15, 1])
                print(f"Found HMC hostname: {hmc_info['hostname']}")
            
            # Model: B3 (row 1, column 1) - Excel row 3 = pandas row 1
            if len(df) > 1 and len(df.columns) > 1:
                hmc_info['hardware_model'] = clean_text(df.iloc[1, 1])
                print(f"Found hardware model: {hmc_info['hardware_model']}")
            
            # Serial: B4 (row 2, column 1) - Excel row 4 = pandas row 2
            if len(df) > 2 and len(df.columns) > 1:
                hmc_info['serial'] = clean_text(df.iloc[2, 1])
                print(f"Found serial: {hmc_info['serial']}")
            
            # Base Version: E6 (row 4, column 4) - Excel row 6 = pandas row 4
            if len(df) > 4 and len(df.columns) > 4:
                hmc_info['base_version'] = clean_text(df.iloc[4, 4])
                print(f"Found base version: {hmc_info['base_version']}")
            
            # Service Pack: E4 (row 2, column 4) - Excel row 4 = pandas row 2
            if len(df) > 2 and len(df.columns) > 4:
                hmc_info['service_pack'] = clean_text(df.iloc[2, 4])
                print(f"Found service pack: {hmc_info['service_pack']}")
            
            # Gateway: B19 (row 17, column 1) - Excel row 19 = pandas row 17
            if len(df) > 17 and len(df.columns) > 1:
                hmc_info['gateway'] = clean_text(df.iloc[17, 1])
                print(f"Found gateway: {hmc_info['gateway']}")
            
            # IP addresses with netmask
            # eth0: B24 (row 22, column 1) + B25 (row 23, column 1) - Excel rows 24,25 = pandas rows 22,23
            if len(df) > 23 and len(df.columns) > 1:
                ip_eth0 = clean_text(df.iloc[22, 1])
                netmask_eth0 = clean_text(df.iloc[23, 1])
                if ip_eth0 and ip_eth0 != 'nan':
                    ip_with_netmask = f"{ip_eth0}/{netmask_eth0}" if netmask_eth0 and netmask_eth0 != 'nan' else ip_eth0
                    hmc_info['ip_addresses']['eth0'] = ip_with_netmask
                    print(f"Found IP for eth0: {ip_with_netmask}")
            
            # eth1: C24 (row 22, column 2) + C25 (row 23, column 2) - Excel rows 24,25 = pandas rows 22,23
            if len(df) > 23 and len(df.columns) > 2:
                ip_eth1 = clean_text(df.iloc[22, 2])
                netmask_eth1 = clean_text(df.iloc[23, 2])
                if ip_eth1 and ip_eth1 != 'nan':
                    ip_with_netmask = f"{ip_eth1}/{netmask_eth1}" if netmask_eth1 and netmask_eth1 != 'nan' else ip_eth1
                    hmc_info['ip_addresses']['eth1'] = ip_with_netmask
                    print(f"Found IP for eth1: {ip_with_netmask}")
            
            # eth2: D24 (row 22, column 3) + D25 (row 23, column 3) - Excel rows 24,25 = pandas rows 22,23
            if len(df) > 23 and len(df.columns) > 3:
                ip_eth2 = clean_text(df.iloc[22, 3])
                netmask_eth2 = clean_text(df.iloc[23, 3])
                if ip_eth2 and ip_eth2 != 'nan':
                    ip_with_netmask = f"{ip_eth2}/{netmask_eth2}" if netmask_eth2 and netmask_eth2 != 'nan' else ip_eth2
                    hmc_info['ip_addresses']['eth2'] = ip_with_netmask
                    print(f"Found IP for eth2: {ip_with_netmask}")
            
            # eth3: E24 (row 22, column 4) + E25 (row 23, column 4) - Excel rows 24,25 = pandas rows 22,23
            if len(df) > 23 and len(df.columns) > 4:
                ip_eth3 = clean_text(df.iloc[22, 4])
                netmask_eth3 = clean_text(df.iloc[23, 4])
                if ip_eth3 and ip_eth3 != 'nan':
                    ip_with_netmask = f"{ip_eth3}/{netmask_eth3}" if netmask_eth3 and netmask_eth3 != 'nan' else ip_eth3
                    hmc_info['ip_addresses']['eth3'] = ip_with_netmask
                    print(f"Found IP for eth3: {ip_with_netmask}")
                    
        except Exception as cell_error:
            print(f"Error accessing specific cells: {cell_error}")
        
        # Add HMC info to data if we found any information
        if any(value for value in hmc_info.values() if value):
            data['hmc_info'].append(hmc_info)
            print(f"Added HMC info: {hmc_info}")
                
    except Exception as e:
        print(f"Error extracting HMC info: {e}")

def extract_system_summary_info_from_sheet(df, data):
    """Extract System Server information from Excel sheet."""
    try:
        print(f"Extracting System Server info from sheet with columns: {list(df.columns)}")
        print(f"DataFrame shape: {df.shape}")
        
        # Extract individual server information from A2 to end of A column
        for index, row in df.iterrows():
            try:
                # Skip header row (index 0) and start from A2 (index 0) - Excel row 2 = pandas row 0
                if index < 0:
                    continue
                
                # Get server name from column A
                server_name = clean_text(row.iloc[0]) if len(row) > 0 else ""
                
                if server_name and server_name != 'nan' and server_name.strip():
                    print(f"Processing server: {server_name}")
                    
                    server_info = {
                        'server_name': server_name,
                        'model': None,
                        'serial': None,
                        'cpu': None,
                        'memory': None,
                        'firmware_level': None,
                        'fsp_ip_address': None
                    }
                    
                    # Server's Model: C cell (column 2)
                    if len(row) > 2:
                        server_info['model'] = clean_text(row.iloc[2])
                        print(f"  Found model: {server_info['model']}")
                    
                    # Server's Serial: D cell (column 3)
                    if len(row) > 3:
                        server_info['serial'] = clean_text(row.iloc[3])
                        print(f"  Found serial: {server_info['serial']}")
                    
                    # Server's CPU: G cell (column 6)
                    if len(row) > 6:
                        server_info['cpu'] = clean_text(row.iloc[6])
                        print(f"  Found CPU: {server_info['cpu']}")
                    
                    # Server's Memory: P cell (column 15)
                    if len(row) > 15:
                        server_info['memory'] = clean_text(row.iloc[15])
                        print(f"  Found memory: {server_info['memory']}")
                    
                    # Server's Firmware Level: AA cell (column 26)
                    if len(row) > 26:
                        server_info['firmware_level'] = clean_text(row.iloc[26])
                        print(f"  Found firmware level: {server_info['firmware_level']}")
                    
                    # Server's FSP IP address: W cell (column 22)
                    if len(row) > 22:
                        server_info['fsp_ip_address'] = clean_text(row.iloc[22])
                        print(f"  Found FSP IP: {server_info['fsp_ip_address']}")
                    
                    # Add server info to data if we found any information
                    if any(value for value in server_info.values() if value and value != server_name):
                        data['system_info'].append(server_info)
                        print(f"Added server info: {server_info}")
                            
            except Exception as row_error:
                print(f"Error processing row {index}: {row_error}")
                continue
                
    except Exception as e:
        print(f"Error extracting System Server info: {e}")

def extract_lpar_info_from_sheet(df, data):
    """Extract LPAR information from Excel sheet."""
    try:
        print(f"Extracting LPAR info from sheet with columns: {list(df.columns)}")
        print(f"DataFrame shape: {df.shape}")
        
        # Extract individual LPAR information from A2 to end of A column
        for index, row in df.iterrows():
            try:
                # Skip header row (index 0) and start from A2 (index 0) - Excel row 2 = pandas row 0
                if index < 0:
                    continue
                
                # Get LPAR name from column A
                lpar_name = clean_text(row.iloc[0]) if len(row) > 0 else ""
                
                if lpar_name and lpar_name != 'nan' and lpar_name.strip():
                    print(f"Processing LPAR: {lpar_name}")
                    
                    lpar_info = {
                        'lpar_name': lpar_name,
                        'desired_entitled_cpu': None,
                        'min_cpu': None,
                        'max_cpu': None,
                        'desired_virtual_processor': None,
                        'min_virtual_processor': None,
                        'max_virtual_processor': None,
                        'entitled_memory_gb': None,
                        'min_memory_gb': None,
                        'max_memory_gb': None,
                        'power_server': None
                    }
                    
                    # LPAR desired entitled CPU: R cell (column 17)
                    if len(row) > 17:
                        lpar_info['desired_entitled_cpu'] = clean_text(row.iloc[17])
                        print(f"  Found desired entitled CPU: {lpar_info['desired_entitled_cpu']}")
                    
                    # LPAR min CPU: Q cell (column 16)
                    if len(row) > 16:
                        lpar_info['min_cpu'] = clean_text(row.iloc[16])
                        print(f"  Found min CPU: {lpar_info['min_cpu']}")
                    
                    # LPAR max CPU: S cell (column 18)
                    if len(row) > 18:
                        lpar_info['max_cpu'] = clean_text(row.iloc[18])
                        print(f"  Found max CPU: {lpar_info['max_cpu']}")
                    
                    # LPAR desired Virtual Processor: U cell (column 20)
                    if len(row) > 20:
                        lpar_info['desired_virtual_processor'] = clean_text(row.iloc[20])
                        print(f"  Found desired virtual processor: {lpar_info['desired_virtual_processor']}")
                    
                    # LPAR min Virtual Processor: T cell (column 19)
                    if len(row) > 19:
                        lpar_info['min_virtual_processor'] = clean_text(row.iloc[19])
                        print(f"  Found min virtual processor: {lpar_info['min_virtual_processor']}")
                    
                    # LPAR max Virtual Processor: V cell (column 21)
                    if len(row) > 21:
                        lpar_info['max_virtual_processor'] = clean_text(row.iloc[21])
                        print(f"  Found max virtual processor: {lpar_info['max_virtual_processor']}")
                    
                    # LPAR entitled memory (GB): H cell (column 7)
                    if len(row) > 7:
                        lpar_info['entitled_memory_gb'] = clean_text(row.iloc[7])
                        print(f"  Found entitled memory (GB): {lpar_info['entitled_memory_gb']}")
                    
                    # LPAR min memory (GB): G cell (column 6)
                    if len(row) > 6:
                        lpar_info['min_memory_gb'] = clean_text(row.iloc[6])
                        print(f"  Found min memory (GB): {lpar_info['min_memory_gb']}")
                    
                    # LPAR max memory (GB): I cell (column 8)
                    if len(row) > 8:
                        lpar_info['max_memory_gb'] = clean_text(row.iloc[8])
                        print(f"  Found max memory (GB): {lpar_info['max_memory_gb']}")
                    
                    # LPAR's Power server: AH cell (column 33)
                    if len(row) > 33:
                        lpar_info['power_server'] = clean_text(row.iloc[33])
                        print(f"  Found power server: {lpar_info['power_server']}")
                    
                    # Add LPAR info to data if we found any information
                    if any(value for value in lpar_info.values() if value and value != lpar_name):
                        data['lpar_info'].append(lpar_info)
                        print(f"Added LPAR info: {lpar_info}")
                            
            except Exception as row_error:
                print(f"Error processing row {index}: {row_error}")
                continue
                
    except Exception as e:
        print(f"Error extracting LPAR info: {e}")

def create_word_document(data, output_file):
    """Create a Word document with the extracted data."""
    doc = Document()
    
    # Title
    title = doc.add_heading('System Server Report', 0)
    title.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Summary
    doc.add_heading('Summary', level=1)
    summary_para = doc.add_paragraph()
    summary_para.add_run(f"Total HMC Systems: {len(data['hmc_info'])}\n")
    summary_para.add_run(f"Total System Servers: {len(data['system_info'])}\n")
    summary_para.add_run(f"Total LPARs: {len(data['lpar_info'])}")
    
    # HMC Information
    if data['hmc_info']:
        doc.add_heading('HMC Information', level=1)
        
        for hmc in data['hmc_info']:
            # Create table for HMC
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add HMC data rows
            if hmc.get('hostname'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'HOSTNAME'
                row_cells[1].text = hmc['hostname']
            
            if hmc.get('hardware_model'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'MODEL'
                row_cells[1].text = hmc['hardware_model']
            
            if hmc.get('serial'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'SERIAL'
                row_cells[1].text = hmc['serial']
            
            if hmc.get('base_version'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'BASE VERSION'
                row_cells[1].text = hmc['base_version']
            
            if hmc.get('service_pack'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'SERVICE PACK'
                row_cells[1].text = hmc['service_pack']
            
            if hmc.get('gateway'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'GATEWAY'
                row_cells[1].text = hmc['gateway']
            
            # Add IP addresses
            if hmc.get('ip_addresses'):
                for interface, ip in hmc['ip_addresses'].items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = f'IP ADDR - {interface.upper()}'
                    row_cells[1].text = ip
            
            doc.add_paragraph()  # Add spacing between tables
    
    # LPAR Information
    if data['lpar_info']:
        doc.add_heading('LPAR Information', level=1)
        
        for lpar in data['lpar_info']:
            # Create table for each LPAR
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add LPAR data rows
            if lpar.get('lpar_name'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'LPAR NAME'
                row_cells[1].text = lpar['lpar_name']
            
            if lpar.get('desired_entitled_cpu'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'DESIRED ENTITLED CPU'
                row_cells[1].text = lpar['desired_entitled_cpu']
            
            if lpar.get('min_cpu'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'MIN CPU'
                row_cells[1].text = lpar['min_cpu']
            
            if lpar.get('max_cpu'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'MAX CPU'
                row_cells[1].text = lpar['max_cpu']
            
            if lpar.get('desired_virtual_processor'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'DESIRED VIRTUAL PROCESSOR'
                row_cells[1].text = lpar['desired_virtual_processor']
            
            if lpar.get('min_virtual_processor'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'MIN VIRTUAL PROCESSOR'
                row_cells[1].text = lpar['min_virtual_processor']
            
            if lpar.get('max_virtual_processor'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'MAX VIRTUAL PROCESSOR'
                row_cells[1].text = lpar['max_virtual_processor']
            
            if lpar.get('entitled_memory_gb'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'ENTITLED MEMORY (GB)'
                row_cells[1].text = lpar['entitled_memory_gb']
            
            if lpar.get('min_memory_gb'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'MIN MEMORY (GB)'
                row_cells[1].text = lpar['min_memory_gb']
            
            if lpar.get('max_memory_gb'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'MAX MEMORY (GB)'
                row_cells[1].text = lpar['max_memory_gb']
            
            if lpar.get('power_server'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'POWER SERVER'
                row_cells[1].text = lpar['power_server']
            
            doc.add_paragraph()  # Add spacing between tables
    
    # System Server Information
    if data['system_info']:
        doc.add_heading('System Server Information', level=1)
        
        for server in data['system_info']:
            # Create table for each server
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add server data rows
            if server.get('server_name'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'SERVER NAME'
                row_cells[1].text = server['server_name']
            
            if server.get('model'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'MODEL'
                row_cells[1].text = server['model']
            
            if server.get('serial'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'SERIAL'
                row_cells[1].text = server['serial']
            
            if server.get('cpu'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'CPU CORES'
                row_cells[1].text = server['cpu']
            
            if server.get('memory'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'MEMORY (GB)'
                row_cells[1].text = server['memory']
            
            if server.get('firmware_level'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'FIRMWARE LEVEL'
                row_cells[1].text = server['firmware_level']
            
            if server.get('fsp_ip_address'):
                row_cells = table.add_row().cells
                row_cells[0].text = 'FSP IP ADDRESS'
                row_cells[1].text = server['fsp_ip_address']
            
            doc.add_paragraph()  # Add spacing between tables
    
    # Save the document
    doc.save(output_file)
    print(f"Report saved to: {output_file}")

def main():
    """Main function to process Excel files and generate report."""
    input_dir = "HMCscannerfile"
    output_file = "System_Server_Report.docx"
    
    # Find all Excel files
    excel_files = glob.glob(os.path.join(input_dir, "*.xls"))
    excel_files.extend(glob.glob(os.path.join(input_dir, "*.xlsx")))
    
    if not excel_files:
        print("No Excel files found in HMCscannerfile directory")
        return
    
    print(f"Found {len(excel_files)} Excel files to process")
    
    # Extract data from all files
    all_data = {
        'hmc_info': [],
        'system_info': [],
        'lpar_info': []
    }
    
    for file_path in excel_files:
        print(f"Processing: {file_path}")
        
        # Extract data from Excel file
        file_data = extract_data_from_excel(file_path)
        
        # Merge data
        for key in all_data:
            all_data[key].extend(file_data[key])
    
    # Generate Word document
    create_word_document(all_data, output_file)
    
    # Print summary
    print("\n" + "="*50)
    print("EXTRACTION SUMMARY")
    print("="*50)
    print(f"Total HMC Systems: {len(all_data['hmc_info'])}")
    print(f"Total System Servers: {len(all_data['system_info'])}")
    print(f"Total LPARs: {len(all_data['lpar_info'])}")

if __name__ == "__main__":
    main()
